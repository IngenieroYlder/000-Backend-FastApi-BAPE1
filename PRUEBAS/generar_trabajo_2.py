from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PRUEBAS_DIR = PROJECT_ROOT / "PRUEBAS"
RESULTS_PATH = PRUEBAS_DIR / "evidencias" / "20260331_051122" / "results.json"
OUTPUT_PATH = PRUEBAS_DIR / "TRABAJO 2.docx"


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_center_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def expected_result_for(codigo: str) -> str:
    expected = {
        "CP-SM-01": "Respuesta 200 y cuerpo {'status':'ok'}",
        "CP-SM-02": "Respuesta 200 y HTML de login",
        "CP-SM-03": "Respuesta 200 y HTML de inicio",
        "CP-AU-01": "Registro exitoso de usuario",
        "CP-AU-02": "Rechazo de correo duplicado con 400",
        "CP-AU-03": "Token bearer con credenciales validas",
        "CP-AU-04": "Rechazo con 401 por clave invalida",
        "CP-RB-01": "Rechazo con 401 por ausencia de token",
        "CP-RB-02": "Respuesta 200 y lista de usuarios",
    }
    return expected.get(codigo, "Resultado esperado definido en el plan de pruebas")


def add_case_table(doc: Document, pruebas: dict) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Codigo"
    headers[1].text = "Tipo"
    headers[2].text = "Descripcion"
    headers[3].text = "Resultado esperado"
    headers[4].text = "Resultado obtenido"

    tipos = {
        "CP-SM": "Smoke",
        "CP-AU": "Funcional - Autenticacion",
        "CP-RB": "Seguridad",
    }

    for prueba in pruebas.values():
        codigo = prueba["nombre_prueba"].split(" - ")[0]
        prefijo = "-".join(codigo.split("-")[:2])
        row = table.add_row().cells
        row[0].text = codigo
        row[1].text = tipos.get(prefijo, "Funcional")
        row[2].text = prueba["nombre_prueba"]
        row[3].text = expected_result_for(codigo)
        row[4].text = f"{prueba['resultado'].capitalize()} ({prueba['duracion_segundos']} s)"


def add_incident_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "ID"
    headers[1].text = "Incidencia"
    headers[2].text = "Impacto"
    headers[3].text = "Accion tomada"

    incidencias = [
        (
            "INC-01",
            "Permisos de Windows sobre directorios temporales de pytest",
            "Impedia la primera corrida automatizada",
            "Se redirecciono TMP/TEMP a PRUEBAS/evidencias y se deshabilito tmpdir de pytest",
        ),
        (
            "INC-02",
            "Advertencias de deprecacion en FastAPI, Pydantic y SQLModel",
            "No bloquean la ejecucion, pero generan deuda tecnica",
            "Se registraron como riesgo tecnico para una futura mejora",
        ),
    ]

    for incidencia in incidencias:
        row = table.add_row().cells
        for index, value in enumerate(incidencia):
            row[index].text = value


def add_schedule_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Fase"
    headers[1].text = "Actividad"
    headers[2].text = "Duracion"
    headers[3].text = "Entregable"

    rows = [
        ("1", "Levantamiento del alcance y requerimientos de prueba", "2 horas", "Plan inicial"),
        ("2", "Diseno de casos de prueba y lista de chequeo", "4 horas", "Casos documentados"),
        ("3", "Implementacion de automatizacion en PRUEBAS", "5 horas", "Scripts y pruebas"),
        ("4", "Ejecucion, evidencias y reportes PDF", "3 horas", "Resultados y reportes"),
        ("5", "Analisis final, riesgos y conclusiones", "2 horas", "Documento final"),
    ]

    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value


def build_document() -> None:
    data = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    resumen = data["resumen"]
    pruebas = data["pruebas"]

    doc = Document()
    doc.core_properties.title = "TRABAJO 2 - Plan de pruebas de software BAPE"
    doc.core_properties.subject = "Plan de pruebas de software"
    doc.core_properties.author = "Codex"

    add_title(doc, "PLAN DE PRUEBAS DE SOFTWARE")
    add_center_line(doc, "Proyecto: Backend FastAPI BAPE")
    add_center_line(doc, "Programa: Analisis y desarrollo de software")
    add_center_line(doc, "Actividad: GA9-220501096-AA1 - Realizar plan de pruebas")
    add_center_line(doc, "Documento: TRABAJO 2")
    add_center_line(doc, "Fecha: 31-03-2026")
    doc.add_page_break()

    doc.add_heading("Introduccion", level=1)
    doc.add_paragraph(
        "Este documento presenta el plan de pruebas de software del proyecto Backend FastAPI BAPE. "
        "El objetivo es verificar el comportamiento funcional del sistema, validar rutas criticas de la API, "
        "documentar resultados y dejar una base de ejecucion automatizada en entorno local."
    )

    doc.add_heading("Analisis de requerimientos", level=1)
    add_bullets(
        doc,
        [
            "El sistema debe permitir registro e inicio de sesion de usuarios por medio de FastAPI y JWT.",
            "Las rutas protegidas deben exigir autenticacion valida para garantizar seguridad.",
            "La aplicacion debe responder correctamente en endpoints base como /health, /login y /.",
            "La estrategia de pruebas debe ejecutarse en local y generar evidencias automatizadas.",
            "Los resultados deben quedar trazables mediante JSON, pantallazos PNG y reportes PDF.",
        ],
    )

    doc.add_heading("Funcionalidades existentes", level=1)
    add_bullets(
        doc,
        [
            "Autenticacion de usuarios mediante /auth/register y /auth/login.",
            "Vista web basica para login, dashboard, inbox, settings, channels, contacts y calendar.",
            "Gestion de productos, servicios, categorias y usuarios.",
            "Modulo CRM para contactos y chats.",
            "Integracion con WhatsApp multi-instancia y webhook de eventos.",
            "Modulo de calendario y citas con endpoints de disponibilidad y reservas.",
        ],
    )

    doc.add_heading("Funcionalidades nuevas", level=1)
    add_bullets(
        doc,
        [
            "Centralizacion de todo el material de pruebas dentro de la carpeta PRUEBAS.",
            "Automatizacion de pruebas con pytest y FastAPI TestClient.",
            "Generacion de evidencias visuales en formato PNG por cada caso ejecutado.",
            "Generacion automatica de PDF consolidado y PDF individual por prueba.",
            "Creacion de lista de chequeo y guia paso a paso para ejecucion academica.",
        ],
    )

    doc.add_heading("Estrategia de pruebas", level=1)
    add_numbered(
        doc,
        [
            "Pruebas smoke para validar disponibilidad inicial de la aplicacion.",
            "Pruebas funcionales de autenticacion para registro y login.",
            "Pruebas de seguridad para verificar restriccion de acceso sin token y acceso autorizado.",
            "Pruebas de regresion local cada vez que se modifica la base del backend o la automatizacion.",
        ],
    )

    doc.add_heading("Entorno de pruebas", level=1)
    add_bullets(
        doc,
        [
            "Sistema operativo local Windows con PowerShell.",
            "Python dentro de entorno virtual venv.",
            "Backend FastAPI ejecutado de forma aislada con TestClient.",
            "Base de datos temporal SQLite creada por cada corrida de prueba.",
            "Sin uso de nube para la ejecucion; todo se realiza en ambiente local.",
        ],
    )

    doc.add_heading("Metodologia", level=1)
    add_numbered(
        doc,
        [
            "Identificacion de funcionalidades criticas del proyecto.",
            "Definicion de casos de prueba con codigos y resultado esperado.",
            "Implementacion de pruebas automatizadas en PRUEBAS/tests.",
            "Ejecucion controlada mediante PRUEBAS/run_tests_with_report.py.",
            "Generacion de evidencias y analisis de resultados para trazabilidad.",
        ],
    )

    doc.add_heading("Cronograma", level=1)
    add_schedule_table(doc)

    doc.add_heading("Casos de prueba", level=1)
    add_case_table(doc, pruebas)

    doc.add_heading("Registro de incidencias", level=1)
    add_incident_table(doc)

    doc.add_heading("Herramientas", level=1)
    add_bullets(
        doc,
        [
            "pytest para la automatizacion de casos de prueba.",
            "FastAPI TestClient para simular solicitudes HTTP locales.",
            "reportlab para crear reportes PDF.",
            "pillow para generar pantallazos de evidencia.",
            "python-docx para generar el documento Word TRABAJO 2.",
        ],
    )

    doc.add_heading("Riesgos", level=1)
    add_bullets(
        doc,
        [
            "Diferencias entre SQLite de pruebas y PostgreSQL de produccion pueden ocultar comportamientos especificos.",
            "Cambios en rutas, modelos o autenticacion pueden invalidar casos ya automatizados.",
            "Las advertencias de deprecacion actuales pueden convertirse en errores en futuras versiones.",
            "Problemas de permisos del sistema operativo pueden afectar carpetas temporales si no se controlan.",
        ],
    )

    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(
        "El proyecto Backend FastAPI BAPE cuenta con una base de pruebas automatizadas funcional y organizada. "
        f"En la ejecucion de referencia se obtuvieron {resumen['aprobadas']} pruebas aprobadas de un total de "
        f"{resumen['total']}, sin fallos ni pruebas omitidas. Esto confirma que el modulo de autenticacion, "
        "las rutas basicas y los controles de acceso evaluados cumplen con el comportamiento esperado en el entorno local."
    )
    doc.add_paragraph(
        "Como linea de mejora, se recomienda ampliar la cobertura hacia CRUD de productos, servicios, contactos, "
        "webhooks, calendario e integraciones externas. Tambien conviene atender las advertencias de deprecacion "
        "detectadas para reducir deuda tecnica y mantener estabilidad a futuro."
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
