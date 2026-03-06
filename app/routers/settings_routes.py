from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlmodel import Session, select
from app.database import get_session
from app import models, auth, schemas
from typing import List, Optional
from pydantic import BaseModel
import shutil
import os
import uuid

router = APIRouter(
    prefix="/api/settings",
    tags=["Settings"]
)

UPLOAD_DIR = "static/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@router.get("/", response_model=models.CompanySettings)
def get_settings(db: Session = Depends(get_session), current_user: models.User = Depends(auth.get_current_user)):
    stmt = select(models.CompanySettings).where(models.CompanySettings.company_id == current_user.company_id)
    settings = db.exec(stmt).first()
    if not settings:
        # Create default if not exists
        settings = models.CompanySettings(company_id=current_user.company_id)
        db.add(settings)
        db.commit()
        db.refresh(settings)
    return settings

@router.post("/update")
def update_settings(
    settings_update: schemas.CompanySettingsUpdate, 
    db: Session = Depends(get_session), 
    current_user: models.User = Depends(auth.get_current_user)
):
    print(f"[Settings] Updating settings for company {current_user.company_id}...")
    stmt = select(models.CompanySettings).where(models.CompanySettings.company_id == current_user.company_id)
    settings = db.exec(stmt).first()
    if not settings:
        print(f"[Settings] No existing settings found for company {current_user.company_id}. Creating new.")
        settings = models.CompanySettings(company_id=current_user.company_id)
        db.add(settings)
        db.commit()
        db.refresh(settings)
    
    # Update fields
    settings_data = settings_update.dict(exclude_unset=True)
    print(f"[Settings] Keys to update: {list(settings_data.keys())}")
    for key, value in settings_data.items():
        if isinstance(value, str) and "api_key" in key:
            value = value.strip()
        setattr(settings, key, value)
        
    db.add(settings)
    db.commit()
    db.refresh(settings)
    
    # Log the update
    try:
        from app.services.buffer_service import buffer_service
        buffer_service._save_log(db, current_user.company_id, "info", "Ajustes", f"⚙️ Configuración guardada: {', '.join(settings_data.keys())}")
    except Exception as e:
        print(f"[Settings] Failed to log update: {e}")

    print("[Settings] Update successful.")
    return settings

@router.post("/upload/branding")
async def upload_branding(
    file: UploadFile = File(...),
    type: str = Form(...), # logo, icon, favicon
    db: Session = Depends(get_session),
    current_user: models.User = Depends(auth.get_current_user)
):
    if type not in ["logo", "icon", "favicon"]:
        raise HTTPException(status_code=400, detail="Invalid upload type")
        
    # Generate filename
    ext = file.filename.split(".")[-1]
    filename = f"{current_user.company_id}_{type}_{uuid.uuid4()}.{ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    # Update URL in DB
    url = f"/{UPLOAD_DIR}/{filename}"
    
    stmt = select(models.CompanySettings).where(models.CompanySettings.company_id == current_user.company_id)
    settings = db.exec(stmt).first()
    if not settings:
        settings = models.CompanySettings(company_id=current_user.company_id)
        db.add(settings)
    
    if type == "logo": settings.logo_url = url
    elif type == "icon": settings.icon_url = url
    elif type == "favicon": settings.favicon_url = url
    
    db.add(settings)
    db.commit()
    
    db.add(settings)
    db.commit()
    
    return {"url": url}

class TestEmailRequest(BaseModel):
    to_email: str

@router.post("/test-email")
async def test_email(
    request: TestEmailRequest,
    db: Session = Depends(get_session),
    current_user: models.User = Depends(auth.get_current_user)
):
    from app.services.email_service import email_service
    
    # Force reload settings to ensure we use what's in DB (or recently saved)
    success = await email_service.send_email(
        to_email=request.to_email,
        subject="Prueba de Configuración SMTP - BAPE",
        body=f"<h1>Hola {current_user.first_name or 'Usuario'},</h1><p>Si estás leyendo esto, la configuración de correo está funcionando correctamente.</p>",
        company_id=current_user.company_id
    )
    
    if success:
        return {"status": "success", "message": "Email enviado correctamente."}
    else:
        raise HTTPException(status_code=400, detail="Error al enviar email. Verifique logs o credenciales.")

# --- DOCUMENT SERVICE HELPERS ---

def extract_text_from_file(file_path: str, ext: str) -> str:
    """
    Extracts text based on file extension.
    """
    text = ""
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text with layout preservation
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                    
                    # Also try to extract tables if text is sparse or structured
                    tables = page.extract_tables()
                    table_text = ""
                    if tables:
                        for table in tables:
                            # Format table as Markdown-ish
                            table_text += "\n"
                            for row in table:
                                # Clean None values and join
                                clean_row = [str(cell) if cell else "" for cell in row]
                                table_text += " | ".join(clean_row) + "\n"
                            table_text += "\n"
                    
                    # Combine existing text with table text if it seems unique, 
                    # but usually extract_text gets it. 
                    # Let's trust extract_text first, it's usually better than pypdf.
                    text += page_text + "\n" + table_text
        
        elif ext in ["doc", "docx"]:
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif ext in ["txt", "md", "csv"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        elif ext in ["jpg", "jpeg", "png"]:
            # Placeholder for Vision API
            text = "[Imagen: Se requiere procesar con IA Vision]"
            
        else:
            text = "[Formato no soportado para lectura de texto]"
            
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        text = f"[Error leyendo archivo: {str(e)}]"
        
    return text.strip()

@router.post("/upload/document")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_session),
    current_user: models.User = Depends(auth.get_current_user)
):
    # Generate filename
    original_name = file.filename
    ext = original_name.split(".")[-1].lower()
    
    # Safe filename with UUID
    safe_filename = f"{current_user.company_id}_doc_{uuid.uuid4()}.{ext}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    # Extract Content
    extracted_text = extract_text_from_file(file_path, ext)
        
    # Create Document Record
    doc = models.KnowledgeDocument(
        company_id=current_user.company_id,
        filename=original_name,
        file_path=file_path,
        content=extracted_text, # Save mined text
        status="processed" if extracted_text else "error"
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # Trigger RAG Processing (Chunking + Embedding) in Background
    from app.services.rag_service import rag_service
    import asyncio
    
    # We can use FastAPI BackgroundTasks, but here we process immediately (async) or add to loop.
    # ideally: background_tasks.add_task(rag_service.process_document, doc.id)
    # Since I didn't inject BackgroundTasks, let's just spawn a task.
    asyncio.create_task(rag_service.process_document(doc.id))

    return doc

@router.get("/documents", response_model=List[models.KnowledgeDocument])
def get_documents(db: Session = Depends(get_session), current_user: models.User = Depends(auth.get_current_user)):
    stmt = select(models.KnowledgeDocument).where(models.KnowledgeDocument.company_id == current_user.company_id)
    return db.exec(stmt).all()

@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_session), current_user: models.User = Depends(auth.get_current_user)):
    doc = db.get(models.KnowledgeDocument, doc_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Document not found")
        
    # Delete file
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
        
    db.delete(doc)
    db.commit()
    return {"status": "deleted"}

@router.get("/documents/{doc_id}")
def get_document_content(doc_id: int, db: Session = Depends(get_session), current_user: models.User = Depends(auth.get_current_user)):
    doc = db.get(models.KnowledgeDocument, doc_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"id": doc.id, "content": doc.content, "filename": doc.filename}

class UpdateDocumentContentRequest(BaseModel):
    content: str

@router.put("/documents/{doc_id}")
def update_document_content(
    doc_id: int, 
    request: UpdateDocumentContentRequest,
    db: Session = Depends(get_session), 
    current_user: models.User = Depends(auth.get_current_user)
):
    doc = db.get(models.KnowledgeDocument, doc_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc.content = request.content
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {"status": "updated", "content": doc.content}

@router.get("/logs", response_model=List[models.SystemLog])
def get_system_logs(
    db: Session = Depends(get_session), 
    current_user: models.User = Depends(auth.get_current_user)
):
    stmt = select(models.SystemLog).where(models.SystemLog.company_id == current_user.company_id).order_by(models.SystemLog.created_at.desc()).limit(100)
    return db.exec(stmt).all()
